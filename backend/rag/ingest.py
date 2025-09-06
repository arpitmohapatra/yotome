"""Document ingestion and processing utilities."""

import os
import logging
import hashlib
import mimetypes
from typing import List, Dict, Any, Optional, Tuple
import uuid
from datetime import datetime
import aiofiles

# Document processing imports
import PyPDF2
from docx import Document as DocxDocument
from bs4 import BeautifulSoup
import markdown

# LangChain text splitters
from langchain.text_splitter import (
    RecursiveCharacterTextSplitter,
    MarkdownHeaderTextSplitter,
    TokenTextSplitter
)

from .settings import settings
from .models import DocumentInfo
from .retriever import vector_store

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Handles document ingestion, chunking, and processing."""
    
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=settings.chunk_size,
            chunk_overlap=settings.chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
        
        self.markdown_splitter = MarkdownHeaderTextSplitter(
            headers_to_split_on=[
                ("#", "Header 1"),
                ("##", "Header 2"),
                ("###", "Header 3"),
            ]
        )
    
    async def process_file(
        self,
        file_path: str,
        filename: str,
        mime_type: str,
        tags: Optional[List[str]] = None
    ) -> Tuple[str, int]:
        """Process a file and add it to the vector store."""
        try:
            # Generate document ID
            doc_id = str(uuid.uuid4())
            
            # Read and extract text from file
            text_content = await self._extract_text(file_path, mime_type)
            
            if not text_content.strip():
                raise ValueError("No text content extracted from file")
            
            # Chunk the text
            chunks = self._chunk_text(text_content, filename, mime_type)
            
            if not chunks:
                raise ValueError("No chunks generated from text content")
            
            # Get file size
            file_size = os.path.getsize(file_path)
            
            # Prepare metadata
            metadata = {
                "filename": filename,
                "mime_type": mime_type,
                "file_size": file_size,
                "tags": tags or [],
                "content_hash": hashlib.md5(text_content.encode()).hexdigest(),
                "processing_timestamp": datetime.utcnow().isoformat()
            }
            
            # Add to vector store
            chunks_added = vector_store.add_document(
                doc_id=doc_id,
                filename=filename,
                chunks=chunks,
                metadata=metadata
            )
            
            logger.info(f"Processed file {filename}: {chunks_added} chunks added")
            return doc_id, chunks_added
            
        except Exception as e:
            logger.error(f"Failed to process file {filename}: {e}")
            raise
    
    async def _extract_text(self, file_path: str, mime_type: str) -> str:
        """Extract text content from a file based on its MIME type."""
        try:
            if mime_type == "application/pdf":
                return await self._extract_pdf_text(file_path)
            elif mime_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return await self._extract_docx_text(file_path)
            elif mime_type in ["text/plain", "text/markdown"]:
                return await self._extract_text_file(file_path)
            elif mime_type == "text/html":
                return await self._extract_html_text(file_path)
            else:
                # Try to read as plain text
                return await self._extract_text_file(file_path)
                
        except Exception as e:
            logger.error(f"Failed to extract text from {file_path}: {e}")
            raise
    
    async def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file."""
        text_content = []
        
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"=== Page {page_num + 1} ===\n{page_text}")
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
        
        return "\n\n".join(text_content)
    
    async def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file."""
        doc = DocxDocument(file_path)
        text_content = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        return "\n\n".join(text_content)
    
    async def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text or markdown file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            return await file.read()
    
    async def _extract_html_text(self, file_path: str) -> str:
        """Extract text from HTML file."""
        async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
            html_content = await file.read()
        
        soup = BeautifulSoup(html_content, 'html.parser')
        
        # Remove script and style elements
        for script in soup(["script", "style"]):
            script.decompose()
        
        # Get text and clean it up
        text = soup.get_text()
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = '\n'.join(chunk for chunk in chunks if chunk)
        
        return text
    
    def _chunk_text(self, text: str, filename: str, mime_type: str) -> List[str]:
        """Chunk text content into smaller pieces."""
        try:
            # For Markdown files, try to split by headers first
            if mime_type == "text/markdown" or filename.endswith('.md'):
                try:
                    md_chunks = self.markdown_splitter.split_text(text)
                    if md_chunks:
                        # Further split large markdown chunks
                        all_chunks = []
                        for chunk in md_chunks:
                            chunk_text = chunk.page_content if hasattr(chunk, 'page_content') else str(chunk)
                            if len(chunk_text) > settings.chunk_size:
                                sub_chunks = self.text_splitter.split_text(chunk_text)
                                all_chunks.extend(sub_chunks)
                            else:
                                all_chunks.append(chunk_text)
                        return all_chunks
                except Exception as e:
                    logger.warning(f"Markdown splitting failed for {filename}: {e}")
            
            # Default recursive splitting
            chunks = self.text_splitter.split_text(text)
            
            # Filter out very short chunks
            filtered_chunks = [
                chunk for chunk in chunks 
                if len(chunk.strip()) > 50  # Minimum chunk length
            ]
            
            return filtered_chunks
            
        except Exception as e:
            logger.error(f"Failed to chunk text for {filename}: {e}")
            raise
    
    def validate_file(self, filename: str, file_size: int, mime_type: str) -> Tuple[bool, str]:
        """Validate uploaded file against constraints."""
        # Check file size
        if file_size > settings.max_file_size:
            return False, f"File size ({file_size} bytes) exceeds maximum allowed ({settings.max_file_size} bytes)"
        
        # Check MIME type
        if mime_type not in settings.allowed_mime_types:
            return False, f"File type '{mime_type}' is not allowed. Allowed types: {', '.join(settings.allowed_mime_types)}"
        
        # Check filename
        if not filename or filename.startswith('.'):
            return False, "Invalid filename"
        
        return True, "File is valid"
    
    def detect_mime_type(self, filename: str, file_path: str) -> str:
        """Detect MIME type of a file."""
        # First try to guess from filename
        mime_type, _ = mimetypes.guess_type(filename)
        
        if mime_type:
            return mime_type
        
        # Try to detect from file content using python-magic
        try:
            import magic
            mime = magic.Magic(mime=True)
            detected_mime = mime.from_file(file_path)
            if detected_mime:
                return detected_mime
        except ImportError:
            logger.warning("python-magic not available for MIME type detection")
        except Exception as e:
            logger.warning(f"Failed to detect MIME type with magic: {e}")
        
        # Fallback based on file extension
        ext = os.path.splitext(filename)[1].lower()
        mime_map = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.txt': 'text/plain',
            '.md': 'text/markdown',
            '.html': 'text/html',
            '.htm': 'text/html'
        }
        
        return mime_map.get(ext, 'application/octet-stream')


# Global document processor instance
document_processor = DocumentProcessor()
